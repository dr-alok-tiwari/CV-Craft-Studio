"""
CV-Craft-Studio - Resume Parser Module
Parses PDF, DOCX, and TXT resumes using local/free libraries only.
No external AI APIs required.
"""

import io
import os
import re
import tempfile
from typing import Dict, List, Optional

MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False


SECTION_HEADINGS = {
    "summary": ["summary", "professional summary", "profile", "objective", "career objective", "about me", "overview", "personal statement", "executive summary", "professional profile"],
    "education": ["education", "educational background", "academic background", "qualifications", "academic qualifications", "degrees", "academics"],
    "experience": ["experience", "work experience", "professional experience", "employment history", "work history", "career history", "employment", "professional background"],
    "internship": ["internship", "internships", "internship experience", "training", "industrial training", "summer training", "apprenticeship"],
    "projects": ["projects", "project experience", "academic projects", "personal projects", "key projects", "notable projects", "project work"],
    "skills": ["skills", "technical skills", "core competencies", "competencies", "key skills", "areas of expertise", "expertise", "technologies", "tools", "programming languages", "software skills"],
    "certifications": ["certifications", "certificates", "professional certifications", "courses", "online courses", "training & certifications", "licenses"],
    "achievements": ["achievements", "accomplishments", "awards", "honors", "recognitions", "honors & awards", "key achievements", "notable achievements"],
    "publications": ["publications", "research papers", "papers", "journals", "conference papers", "research publications", "articles"],
    "leadership": ["leadership", "positions of responsibility", "extracurricular activities", "activities", "volunteer", "volunteering", "community service", "co-curricular", "extra-curricular"],
    "languages": ["languages", "language skills", "language proficiency"],
    "references": ["references", "referees"],
}

HEADING_LOOKUP: Dict[str, str] = {
    heading.lower(): key
    for key, headings in SECTION_HEADINGS.items()
    for heading in headings
}


def _validate_file_size(file_bytes: bytes, filename: str = "uploaded file") -> None:
    size = len(file_bytes or b"")
    if size > MAX_UPLOAD_BYTES:
        mb = size / (1024 * 1024)
        raise ValueError(f"{filename} is {mb:.1f} MB. Please upload a resume under 10 MB.")


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber, with PyPDF2 fallback."""
    text = ""
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text += page_text + "\n"
            if text.strip():
                return text
        except Exception:
            pass

    if PYPDF2_AVAILABLE:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text += page_text + "\n"
        except Exception:
            pass
    return text


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes, including simple table text."""
    text = ""
    if DOCX_AVAILABLE:
        try:
            document = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            table_rows = []
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_rows.append(row_text)
            text = "\n".join(paragraphs + table_rows)
            if text.strip():
                return text
        except Exception:
            pass

    if DOCX2TXT_AVAILABLE:
        temp_path = ""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file_bytes)
                temp_path = tmp.name
            text = docx2txt.process(temp_path) or ""
        except Exception:
            pass
        finally:
            if temp_path and os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except OSError:
                    pass
    return text


def parse_txt(file_bytes: bytes) -> str:
    """Decode plain text using common encodings."""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def parse_resume_file(file_bytes: bytes, filename: str) -> str:
    """Dispatch parser based on file extension after enforcing the 10 MB limit."""
    _validate_file_size(file_bytes, filename)
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    if ext == "pdf":
        return parse_pdf(file_bytes)
    if ext in {"docx", "doc"}:
        return parse_docx(file_bytes)
    if ext == "txt":
        return parse_txt(file_bytes)
    return parse_txt(file_bytes)


def extract_email(text: str) -> str:
    match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text or "")
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    patterns = [
        r"(?:\+91[\-\s]?)?(?:\(?\d{3,5}\)?[\-\s]?)?\d{3}[\-\s]?\d{4}",
        r"(?:\+1[\s\-]?)?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4}",
        r"\+?\d{1,3}[\-\s]?\d{4,5}[\-\s]?\d{4,5}",
    ]
    for pattern in patterns:
        match = re.search(pattern, text or "")
        if match:
            return match.group(0).strip()
    return ""


def extract_linkedin(text: str) -> str:
    match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9\-_%]+)", text or "", re.I)
    if match:
        return f"linkedin.com/in/{match.group(1)}"
    match = re.search(r"linkedin:\s*([a-zA-Z0-9\-_%]+)", text or "", re.I)
    return f"linkedin.com/in/{match.group(1)}" if match else ""


def extract_github(text: str) -> str:
    match = re.search(r"(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9\-_]+)", text or "", re.I)
    return f"github.com/{match.group(1)}" if match else ""


def extract_name(text: str) -> str:
    """Heuristic: first non-empty header line that looks like a person name."""
    lines = [line.strip() for line in (text or "").split("\n") if line.strip()]
    for line in lines[:8]:
        if re.search(r"@|\d{5,}|linkedin|github|http|www\.", line, re.I):
            continue
        clean = re.sub(r"[^A-Za-z .'-]", "", line).strip()
        words = clean.split()
        if 2 <= len(words) <= 5 and sum(1 for w in words if w[:1].isupper()) >= min(2, len(words)):
            return clean
    return lines[0] if lines else ""


def extract_contact_info(text: str) -> Dict[str, str]:
    return {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "linkedin": extract_linkedin(text),
        "github": extract_github(text),
    }


def _normalise_heading(line: str) -> str:
    clean = line.strip().strip("•:-—–|").strip()
    clean = re.sub(r"\s+", " ", clean)
    clean = re.sub(r"^\d+[.)]\s*", "", clean)
    return clean.lower()


def _is_section_heading(line: str) -> Optional[str]:
    clean = _normalise_heading(line)
    if not clean or len(clean.split()) > 7:
        return None
    return HEADING_LOOKUP.get(clean)


def extract_sections(text: str) -> Dict[str, str]:
    """Split resume text into sections using heading detection."""
    sections: Dict[str, List[str]] = {"_header": []}
    current_section = "_header"
    for raw_line in (text or "").split("\n"):
        section_key = _is_section_heading(raw_line)
        if section_key:
            current_section = section_key
            sections.setdefault(current_section, [])
            continue
        sections.setdefault(current_section, []).append(raw_line)
    return {key: "\n".join(value).strip() for key, value in sections.items()}


STOPWORDS = {
    "a", "an", "the", "and", "or", "but", "in", "on", "at", "to", "for", "of", "with", "by", "from",
    "is", "are", "was", "were", "be", "been", "being", "have", "has", "had", "do", "does", "did", "will",
    "would", "could", "should", "may", "might", "shall", "can", "i", "me", "my", "we", "our", "you", "your",
    "he", "she", "it", "they", "their", "this", "that", "these", "those", "as", "if", "not", "no", "so", "up",
    "out", "about", "than", "more", "into", "also", "its", "well", "one", "two",
}


def extract_keywords(text: str, top_n: int = 30) -> List[str]:
    """Extract top keywords from text using TF-IDF with frequency fallback."""
    text = text or ""
    if not text.strip():
        return []
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2), max_features=200, token_pattern=r"[a-zA-Z][a-zA-Z0-9+#.]*[a-zA-Z0-9]")
        tfidf_matrix = vectorizer.fit_transform([text])
        scores = zip(vectorizer.get_feature_names_out(), tfidf_matrix.toarray()[0])
        return [word for word, score in sorted(scores, key=lambda item: item[1], reverse=True)[:top_n] if score > 0]
    except Exception:
        words = re.findall(r"[a-zA-Z][a-zA-Z0-9+#.]*", text.lower())
        freq: Dict[str, int] = {}
        for word in words:
            if word not in STOPWORDS and len(word) > 2:
                freq[word] = freq.get(word, 0) + 1
        return [word for word, _ in sorted(freq.items(), key=lambda item: item[1], reverse=True)[:top_n]]


def parse_resume(file_bytes: bytes, filename: str) -> Dict:
    """Full pipeline: file bytes to structured resume dictionary."""
    raw_text = parse_resume_file(file_bytes, filename)
    contact = extract_contact_info(raw_text)
    sections = extract_sections(raw_text)
    keywords = extract_keywords(raw_text)
    return {
        "raw_text": raw_text,
        "contact": contact,
        "sections": sections,
        "keywords": keywords,
        "filename": filename,
        "char_count": len(raw_text),
        "word_count": len(raw_text.split()),
        "line_count": len(raw_text.split("\n")),
    }
