"""
CV-Craft-Studio - Export Utilities
Export resume data to PDF, DOCX, and TXT using local libraries only.

`resume_data_to_text` returns a normal string subclass that also supports
`.decode()`. This prevents older app code from crashing when it calls decode()
on text generated from the builder.
"""

import io
import re
import html
from typing import Dict, List
from datetime import datetime

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DecodableText(str):
    """A string with a decode method for backward compatibility."""
    def decode(self, encoding: str = "utf-8", errors: str = "strict") -> str:
        return str(self)


def _items(value) -> List[str]:
    if not value:
        return []
    if isinstance(value, (list, tuple, set)):
        return [str(v).strip() for v in value if str(v).strip()]
    return [str(value).strip()] if str(value).strip() else []


def _section(lines: List[str], title: str) -> None:
    lines.append(title.upper())
    lines.append("-" * len(title))


def _bullets(lines: List[str], values) -> None:
    for value in values or []:
        clean = str(value).strip().lstrip("•-* ").strip()
        if clean:
            lines.append(f"  • {clean}")


def resume_data_to_text(resume_data: Dict) -> str:
    """Convert resume data dict to clean ATS-readable plain text."""
    lines: List[str] = []
    p = resume_data.get("personal", {}) or {}

    name = str(p.get("name", "")).strip().upper()
    if name:
        lines.extend([name, "=" * len(name)])

    contact = [p.get("email"), p.get("phone"), p.get("linkedin"), p.get("github")]
    contact = [str(x).strip() for x in contact if str(x).strip()]
    if contact:
        lines.append(" | ".join(contact))
    if p.get("location"):
        lines.append(str(p.get("location")).strip())
    lines.append("")

    if resume_data.get("summary"):
        _section(lines, "Professional Summary")
        lines.append(str(resume_data.get("summary", "")).strip())
        lines.append("")

    if resume_data.get("education"):
        _section(lines, "Education")
        for edu in resume_data.get("education", []) or []:
            parts = [edu.get("degree"), edu.get("institution"), edu.get("year")]
            lines.append(" | ".join(str(x).strip() for x in parts if str(x).strip()))
            for field in ("grade", "details"):
                if edu.get(field):
                    lines.append(f"  {edu[field]}")
        lines.append("")

    if resume_data.get("experience"):
        _section(lines, "Professional Experience")
        for exp in resume_data.get("experience", []) or []:
            parts = [exp.get("title"), exp.get("company"), exp.get("duration")]
            lines.append(" | ".join(str(x).strip() for x in parts if str(x).strip()))
            if exp.get("location"):
                lines.append(f"  {exp.get('location')}")
            _bullets(lines, exp.get("bullets", []))
        lines.append("")

    if resume_data.get("internships"):
        _section(lines, "Internships")
        for exp in resume_data.get("internships", []) or []:
            parts = [exp.get("title"), exp.get("company"), exp.get("duration")]
            lines.append(" | ".join(str(x).strip() for x in parts if str(x).strip()))
            _bullets(lines, exp.get("bullets", []))
        lines.append("")

    if resume_data.get("projects"):
        _section(lines, "Projects")
        for proj in resume_data.get("projects", []) or []:
            parts = [proj.get("title"), proj.get("tech"), proj.get("year")]
            lines.append(" | ".join(str(x).strip() for x in parts if str(x).strip()))
            _bullets(lines, proj.get("bullets", []))
        lines.append("")

    if resume_data.get("skills"):
        _section(lines, "Skills")
        skills = resume_data.get("skills")
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                skill_items = _items(skill_list)
                if skill_items:
                    lines.append(f"  {str(category).title()}: {', '.join(skill_items)}")
        else:
            skill_items = _items(skills)
            if skill_items:
                lines.append("  " + ", ".join(skill_items))
        lines.append("")

    if resume_data.get("certifications"):
        _section(lines, "Certifications")
        for cert in resume_data.get("certifications", []) or []:
            if isinstance(cert, dict):
                parts = [cert.get("name"), cert.get("issuer"), cert.get("year")]
                line = " | ".join(str(x).strip() for x in parts if str(x).strip())
            else:
                line = str(cert).strip()
            if line:
                lines.append(f"  • {line}")
        lines.append("")

    if resume_data.get("achievements"):
        _section(lines, "Achievements")
        _bullets(lines, resume_data.get("achievements", []))
        lines.append("")

    if resume_data.get("publications"):
        _section(lines, "Publications")
        for i, pub in enumerate(resume_data.get("publications", []) or [], 1):
            if isinstance(pub, dict):
                parts = [pub.get("title"), pub.get("journal"), pub.get("year")]
                line = " — ".join(str(x).strip() for x in parts if str(x).strip())
            else:
                line = str(pub).strip()
            if line:
                lines.append(f"  {i}. {line}")
        lines.append("")

    if resume_data.get("languages"):
        _section(lines, "Languages")
        lines.append("  " + ", ".join(_items(resume_data.get("languages"))))
        lines.append("")

    return DecodableText("\n".join(lines).strip() + "\n")


def export_txt(resume_data: Dict) -> bytes:
    """Export resume as UTF-8 plain text bytes."""
    return str(resume_data_to_text(resume_data)).encode("utf-8")


def _p(text) -> str:
    return html.escape(str(text or ""), quote=False).replace("\n", "<br/>")


def export_pdf(resume_data: Dict, template_name: str = "ats_classic") -> bytes:
    """Export resume as a clean single-column ATS-friendly PDF."""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("reportlab is not installed. Run: pip install reportlab")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=1.35*cm, leftMargin=1.35*cm, topMargin=1.25*cm, bottomMargin=1.25*cm)
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("ResumeName", parent=styles["Title"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=colors.HexColor("#1a1a2e"), spaceAfter=3)
    contact_style = ParagraphStyle("ResumeContact", parent=styles["Normal"], alignment=TA_CENTER, fontName="Helvetica", fontSize=8.5, leading=11, textColor=colors.HexColor("#555555"), spaceAfter=8)
    section_style = ParagraphStyle("ResumeSection", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=colors.HexColor("#1a1a2e"), spaceBefore=8, spaceAfter=2)
    body_style = ParagraphStyle("ResumeBody", parent=styles["Normal"], fontName="Helvetica", fontSize=9.2, leading=12.5, textColor=colors.HexColor("#222222"), spaceAfter=2)
    bullet_style = ParagraphStyle("ResumeBullet", parent=body_style, leftIndent=12, firstLineIndent=-8)
    small_style = ParagraphStyle("ResumeSmall", parent=body_style, fontSize=8.8, textColor=colors.HexColor("#555555"))
    story = []
    p = resume_data.get("personal", {}) or {}
    story.append(Paragraph(_p(str(p.get("name") or "Your Name").upper()), name_style))
    contact = [p.get("email"), p.get("phone"), p.get("linkedin"), p.get("github"), p.get("location")]
    contact = [str(x).strip() for x in contact if str(x).strip()]
    if contact:
        story.append(Paragraph(_p(" | ".join(contact)), contact_style))
    story.append(HRFlowable(width="100%", thickness=1.1, color=colors.HexColor("#1a1a2e")))

    def add_section(title: str) -> None:
        story.append(Paragraph(_p(title.upper()), section_style))
        story.append(HRFlowable(width="100%", thickness=0.45, color=colors.HexColor("#D1D5DB")))

    def add_text(text, style=body_style) -> None:
        if str(text or "").strip():
            story.append(Paragraph(_p(text), style))

    def add_rich_text(text, style=body_style) -> None:
        if str(text or "").strip():
            story.append(Paragraph(str(text), style))

    def add_bullet(text) -> None:
        clean = str(text or "").strip().lstrip("•-* ").strip()
        if clean:
            story.append(Paragraph(f"• {_p(clean)}", bullet_style))

    if resume_data.get("summary"):
        add_section("Professional Summary"); add_text(resume_data.get("summary"))
    if resume_data.get("education"):
        add_section("Education")
        for edu in resume_data.get("education", []) or []:
            line = " | ".join(str(x).strip() for x in [edu.get("degree"), edu.get("institution"), edu.get("year")] if str(x).strip())
            add_rich_text(f"<b>{_p(line)}</b>")
            for field in ("grade", "details"):
                if edu.get(field): add_text(edu.get(field), small_style)
            story.append(Spacer(1, 3))
    if resume_data.get("experience"):
        add_section("Professional Experience")
        for exp in resume_data.get("experience", []) or []:
            line = " | ".join(str(x).strip() for x in [exp.get("title"), exp.get("company"), exp.get("duration")] if str(x).strip())
            add_rich_text(f"<b>{_p(line)}</b>")
            if exp.get("location"): add_text(exp.get("location"), small_style)
            for b in exp.get("bullets", []) or []: add_bullet(b)
            story.append(Spacer(1, 4))
    if resume_data.get("internships"):
        add_section("Internships")
        for exp in resume_data.get("internships", []) or []:
            line = " | ".join(str(x).strip() for x in [exp.get("title"), exp.get("company"), exp.get("duration")] if str(x).strip())
            add_rich_text(f"<b>{_p(line)}</b>")
            for b in exp.get("bullets", []) or []: add_bullet(b)
            story.append(Spacer(1, 3))
    if resume_data.get("projects"):
        add_section("Projects")
        for proj in resume_data.get("projects", []) or []:
            line = " | ".join(str(x).strip() for x in [proj.get("title"), proj.get("tech"), proj.get("year")] if str(x).strip())
            add_rich_text(f"<b>{_p(line)}</b>")
            for b in proj.get("bullets", []) or []: add_bullet(b)
            story.append(Spacer(1, 3))
    if resume_data.get("skills"):
        add_section("Skills")
        skills = resume_data.get("skills")
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                skill_items = _items(skill_list)
                if skill_items: add_rich_text(f"<b>{_p(str(category).title())}:</b> {_p(', '.join(skill_items))}")
        else:
            add_text(", ".join(_items(skills)))
    if resume_data.get("certifications"):
        add_section("Certifications")
        for cert in resume_data.get("certifications", []) or []:
            line = " | ".join(str(x).strip() for x in [cert.get("name"), cert.get("issuer"), cert.get("year")] if str(x).strip()) if isinstance(cert, dict) else str(cert).strip()
            add_bullet(line)
    if resume_data.get("achievements"):
        add_section("Achievements")
        for ach in resume_data.get("achievements", []) or []: add_bullet(ach)
    if resume_data.get("publications"):
        add_section("Publications")
        for i, pub in enumerate(resume_data.get("publications", []) or [], 1):
            line = " — ".join(str(x).strip() for x in [pub.get("title"), pub.get("journal"), pub.get("year")] if str(x).strip()) if isinstance(pub, dict) else str(pub).strip()
            add_text(f"{i}. {line}")
    if resume_data.get("languages"):
        add_section("Languages"); add_text(", ".join(_items(resume_data.get("languages"))))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def export_docx(resume_data: Dict, template_name: str = "ats_classic") -> bytes:
    """Export resume as a clean ATS-friendly DOCX."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.4); section.bottom_margin = Cm(1.4); section.left_margin = Cm(1.4); section.right_margin = Cm(1.4)
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    p = resume_data.get("personal", {}) or {}
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(str(p.get("name") or "Your Name").upper()); run.bold = True; run.font.size = Pt(16)
    contact = [p.get("email"), p.get("phone"), p.get("linkedin"), p.get("github"), p.get("location")]
    contact = [str(x).strip() for x in contact if str(x).strip()]
    if contact:
        cp = doc.add_paragraph(" | ".join(contact)); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def heading(text):
        h = doc.add_paragraph(); r = h.add_run(text.upper()); r.bold = True; r.font.size = Pt(11)
    def para(text, bold=False):
        if str(text or "").strip():
            pg = doc.add_paragraph(); r = pg.add_run(str(text).strip()); r.bold = bold
    def bullet(text):
        clean = str(text or "").strip().lstrip("•-* ").strip()
        if clean: doc.add_paragraph(clean, style="List Bullet")

    if resume_data.get("summary"):
        heading("Professional Summary"); para(resume_data.get("summary"))
    for key, title_name in [("education", "Education"), ("experience", "Professional Experience"), ("internships", "Internships"), ("projects", "Projects")]:
        if resume_data.get(key):
            heading(title_name)
            for item in resume_data.get(key, []) or []:
                if key == "education":
                    line = " | ".join(str(x).strip() for x in [item.get("degree"), item.get("institution"), item.get("year")] if str(x).strip())
                    para(line, True)
                    for field in ("grade", "details"):
                        if item.get(field): para(item.get(field))
                elif key == "projects":
                    line = " | ".join(str(x).strip() for x in [item.get("title"), item.get("tech"), item.get("year")] if str(x).strip())
                    para(line, True)
                    for b in item.get("bullets", []) or []: bullet(b)
                else:
                    line = " | ".join(str(x).strip() for x in [item.get("title"), item.get("company"), item.get("duration")] if str(x).strip())
                    para(line, True)
                    if item.get("location"): para(item.get("location"))
                    for b in item.get("bullets", []) or []: bullet(b)
    if resume_data.get("skills"):
        heading("Skills")
        skills = resume_data.get("skills")
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                skill_items = _items(skill_list)
                if skill_items: para(f"{str(category).title()}: {', '.join(skill_items)}")
        else:
            para(", ".join(_items(skills)))
    if resume_data.get("certifications"):
        heading("Certifications")
        for cert in resume_data.get("certifications", []) or []:
            line = " | ".join(str(x).strip() for x in [cert.get("name"), cert.get("issuer"), cert.get("year")] if str(x).strip()) if isinstance(cert, dict) else str(cert).strip()
            bullet(line)
    if resume_data.get("achievements"):
        heading("Achievements")
        for ach in resume_data.get("achievements", []) or []: bullet(ach)
    if resume_data.get("publications"):
        heading("Publications")
        for i, pub in enumerate(resume_data.get("publications", []) or [], 1):
            line = " — ".join(str(x).strip() for x in [pub.get("title"), pub.get("journal"), pub.get("year")] if str(x).strip()) if isinstance(pub, dict) else str(pub).strip()
            para(f"{i}. {line}")
    if resume_data.get("languages"):
        heading("Languages"); para(", ".join(_items(resume_data.get("languages"))))

    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer.read()


def get_export_filename(resume_data: Dict, extension: str) -> str:
    """Generate a clean timestamped export filename."""
    p = resume_data.get("personal", {}) or {}
    name = str(p.get("name") or "resume").strip().lower()
    name = re.sub(r"[^a-z0-9]+", "_", name).strip("_") or "resume"
    ext = extension.lower().lstrip(".")
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    return f"{name}_ats_ready_{stamp}.{ext}"
